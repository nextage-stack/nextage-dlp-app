"""Build Nextage_DLP_Azure_Deployment.docx from the deployment guide.

Generates a self-contained Word document covering the IT admin steps to
deploy the add-in to Azure, plus a short end-user section showing what the
add-in looks like in Outlook once deployed.
"""

from __future__ import annotations

import os
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Cm, Inches

REPO = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT = os.path.join(REPO, "Nextage_DLP_Azure_Deployment.docx")


# --------------------------------------------------------------------------- #
# Helpers
# --------------------------------------------------------------------------- #


def set_cell_shading(cell, color_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def add_horizontal_rule(doc: Document) -> None:
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "BFBFBF")
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_code_block(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = "Menlo"
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x1B, 0x1B, 0x1B)
    # Light grey background via paragraph shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F2F2F2")
    pPr.append(shd)


def add_callout(doc: Document, title: str, body: str, color: str = "FFF4CE") -> None:
    """One-row, one-cell highlight box."""
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.cell(0, 0)
    set_cell_shading(cell, color)
    cell.text = ""
    para = cell.paragraphs[0]
    bold = para.add_run(f"{title}  ")
    bold.bold = True
    para.add_run(body)
    doc.add_paragraph()


def add_bullets(doc: Document, items: list[str]) -> None:
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def add_check_items(doc: Document, items: list[str]) -> None:
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run("☐  ")
        run.font.name = "Menlo"
        p.add_run(it)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.autofit = True
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = True
        set_cell_shading(hdr[i], "DEEBF7")
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = ""
            cell.paragraphs[0].add_run(val)
    doc.add_paragraph()


# --------------------------------------------------------------------------- #
# Document content
# --------------------------------------------------------------------------- #


def build() -> None:
    doc = Document()

    # Base style
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    # ---------------- Title ----------------
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = title.add_run("Nextage DLP — Azure Deployment & Outlook Rollout")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = RGBColor(0x35, 0x25, 0xCD)

    subtitle = doc.add_paragraph()
    sr = subtitle.add_run("How to publish the DLP add-in to Azure and make it usable inside an end-user's Outlook")
    sr.italic = True
    sr.font.size = Pt(12)
    sr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    add_horizontal_rule(doc)

    # ---------------- Overview ----------------
    doc.add_heading("1. Overview", level=1)
    doc.add_paragraph(
        "Nextage DLP is an Outlook add-in that runs three security checks on every "
        "outgoing email (attachment encryption, filename–client match, subject + "
        "recipient-domain validation). The add-in itself is delivered as static HTML/JS "
        "from Azure Static Web Apps, talks to a small Azure Functions API for "
        "configuration and audit logging, and reads its DLP rules from Cosmos DB."
    )
    doc.add_paragraph(
        "This document walks through a full first-time deployment. Allow 60–90 minutes "
        "end-to-end. Every phase ends with verification checks you should tick off "
        "before moving on. Where two paths exist, the recommended one is marked."
    )

    doc.add_heading("Architecture at a glance", level=2)
    add_bullets(doc, [
        "Static Web App  →  hosts taskpane.html / commands.js (the add-in UI).",
        "Azure Functions  →  /api/config and /api/audit, behind Entra ID JWT validation.",
        "Cosmos DB (Serverless)  →  customers, advisors, exemptions, exclusions, audit log.",
        "Entra ID App Registration  →  Office.js SSO + delegated API access.",
        "Azure DevOps pipeline  →  builds the front-end + API, substitutes manifest tokens, deploys with a manual approval gate.",
    ])

    add_callout(
        doc,
        "Heads up:",
        "Cosmos master-key access is disabled by design. All data-plane access is via "
        "Managed Identity. If you need to seed or edit Cosmos directly, you must use AAD "
        "(documented in Phase 5).",
    )

    # ---------------- Phase 0: Prerequisites ----------------
    doc.add_heading("2. Phase 0 — Tooling and identities", level=1)
    doc.add_paragraph("Install on your workstation:")
    add_bullets(doc, [
        "Azure CLI ≥ 2.55  (brew install azure-cli)",
        "Terraform ≥ 1.7  (brew install terraform)",
        "Node.js 20.x  (nvm install 20 && nvm use 20)",
        "jq  (brew install jq)",
    ])

    doc.add_paragraph("Roles you (or a colleague who helps you) need:")
    add_table(
        doc,
        ["Role", "Where", "Used for"],
        [
            ["Owner or Contributor + User Access Administrator", "Azure subscription", "Create resources, assign RBAC"],
            ["Application Administrator or Global Administrator", "Entra ID tenant", "Create the app registration, grant admin consent"],
            ["Global Administrator", "Microsoft 365 tenant", "Upload the add-in to Integrated apps, assign to users"],
        ],
    )

    doc.add_paragraph("Sign in to Azure once:")
    add_code_block(doc, "az login\naz account set --subscription <SUBSCRIPTION_ID>\naz account show")

    add_callout(
        doc,
        "Verify:",
        "az account show lists the correct subscription and tenant before continuing.",
    )

    # ---------------- Phase 1: Persistent secrets ----------------
    doc.add_heading("3. Phase 1 — Generate the two persistent secrets", level=1)
    doc.add_paragraph(
        "Both secrets below are generated ONCE and never rotated under normal "
        "operations. The manifest GUID is the add-in's identity to Office — if you "
        "regenerate it, every existing install treats the deploy as a brand-new "
        "add-in. The HMAC key hashes PII (subject, attachment names, recipients) in "
        "audit log entries; rotating it invalidates historical hashes."
    )

    doc.add_paragraph("Generate the manifest GUID:")
    add_code_block(doc, "uuidgen | tr '[:upper:]' '[:lower:]'\n# → e.g. 8f4a2d1c-9b3e-4ad5-bc62-3f1c2d4e5f6a")

    doc.add_paragraph("Generate the audit HMAC key:")
    add_code_block(doc, "openssl rand -hex 32\n# → e.g. 7b8c9d... (64 hex chars)")

    add_callout(
        doc,
        "Store safely:",
        "Save both values in a password manager or Key Vault you control. You will "
        "paste them into terraform.tfvars in the next phase.",
    )

    # ---------------- Phase 2: Terraform ----------------
    doc.add_heading("4. Phase 2 — Provision Azure infrastructure", level=1)
    doc.add_paragraph("From the project root:")
    add_code_block(doc, "cd terraform\ncp terraform.tfvars.example terraform.tfvars")

    doc.add_paragraph("Edit terraform.tfvars and fill in the required values:")
    add_code_block(doc, (
        "subscription_id = \"<your subscription id>\"\n"
        "tenant_id       = \"<your tenant id>\"\n"
        "manifest_guid   = \"<UUID from Phase 1>\"\n"
        "audit_hmac_key  = \"<hex string from Phase 1>\"\n"
        "\n"
        "# Optional — your office IP for Cosmos break-glass admin\n"
        "# cosmos_admin_ips = [\"203.0.113.0/32\"]"
    ))

    doc.add_paragraph("Run Terraform:")
    add_code_block(doc, "terraform init\nterraform plan      # read the plan carefully\nterraform apply     # type 'yes'")

    doc.add_paragraph("When apply completes, capture every output you'll need later:")
    add_code_block(doc, "terraform output -json > /tmp/dlp-outputs.json\nterraform output deployment_summary")

    add_check_items(doc, [
        "terraform apply finished with no errors.",
        "terraform output entra_app_client_id returns a GUID.",
        "terraform output static_web_app_hostname returns something like nextage-dlp-addin-12345.azurestaticapps.net.",
    ])

    # ---------------- Phase 3: Admin consent ----------------
    doc.add_heading("5. Phase 3 — Grant admin consent for the Entra app", level=1)
    doc.add_paragraph(
        "Until admin consent is granted, end users will get a consent prompt on first "
        "use (and may not be able to grant it themselves). A Global or Application "
        "Administrator runs:"
    )
    add_code_block(doc, (
        "CLIENT_ID=$(terraform output -raw entra_app_client_id)\n"
        "az ad app permission admin-consent --id \"$CLIENT_ID\""
    ))

    doc.add_paragraph("Then verify in the Entra Admin Center:")
    add_bullets(doc, [
        "Open Entra Admin Center  →  App registrations  →  Nextage DLP Add-in.",
        "API permissions  →  every row shows 'Granted for <Tenant>' with a green check.",
        "Expose an API  →  scope 'access_as_user' exists; Application ID URI matches terraform output entra_app_identifier_uri exactly.",
    ])

    add_check_items(doc, [
        "Admin consent granted (no yellow warning banners).",
        "Application ID URI matches the Terraform output.",
    ])

    # ---------------- Phase 4: Pipeline ----------------
    doc.add_heading("6. Phase 4 — Configure the Azure DevOps pipeline", level=1)

    doc.add_heading("6.1 Variable group", level=2)
    doc.add_paragraph(
        "Azure DevOps  →  Pipelines  →  Library  →  + Variable group  →  name it "
        "exactly 'nextage-dlp-prod'. Add these variables (use terraform output to "
        "look up each value):"
    )

    add_table(
        doc,
        ["Variable name", "Source", "Mark as secret?"],
        [
            ["AZURE_STATIC_WEB_APPS_API_TOKEN", "terraform output -raw static_web_app_deployment_token", "Yes"],
            ["AZURE_FUNCTIONS_URL", "terraform output -raw function_app_url", "No"],
            ["SWA_HOSTNAME", "terraform output -json manifest_tokens | jq -r .SWA_HOSTNAME", "No"],
            ["FUNCTIONS_HOSTNAME", "terraform output -json manifest_tokens | jq -r .FUNCTIONS_HOSTNAME", "No"],
            ["ENTRA_APP_CLIENT_ID", "terraform output -raw entra_app_client_id", "No"],
            ["MANIFEST_GUID", "The UUID from Phase 1", "No"],
            ["FUNCTION_APP_NAME", "terraform output -raw function_app_name", "No"],
        ],
    )

    doc.add_heading("6.2 Service connection", level=2)
    doc.add_paragraph(
        "Project Settings  →  Service connections  →  New  →  Azure Resource Manager  "
        "→  Workload Identity Federation (recommended). Scope it to the resource "
        "group. Name it exactly 'nextage-azure-prod' (the pipeline references this string)."
    )

    doc.add_heading("6.3 Approval gates", level=2)
    doc.add_paragraph(
        "Pipelines  →  Environments  →  New, for each of:"
    )
    add_bullets(doc, ["production-addin", "production-api"])
    doc.add_paragraph(
        "On each: Approvals and checks  →  +  →  Approvals. Add yourself (and ideally "
        "one other person) as a required approver. Without this gate, every merge to "
        "main auto-deploys to production."
    )

    add_check_items(doc, [
        "Variable group has all 7 values, with the API token marked as a secret.",
        "Service connection nextage-azure-prod created.",
        "Both environments have manual approval gates.",
    ])

    # ---------------- Phase 5: Seed Cosmos ----------------
    doc.add_heading("7. Phase 5 — Seed Cosmos DB with initial DLP rules", level=1)
    doc.add_paragraph(
        "The add-in does nothing useful until you populate the customer / advisor / "
        "exemption / exclusion lists. Cosmos local-auth is off, so seeding is done "
        "via AAD with a short-lived role grant to your user."
    )

    doc.add_heading("7.1 Grant yourself temporary Cosmos access", level=2)
    add_code_block(doc, (
        "COSMOS_ACCT=$(terraform output -raw cosmos_account_name)\n"
        "RG=$(terraform output -raw resource_group_name)\n"
        "MY_OID=$(az ad signed-in-user show --query id -o tsv)\n"
        "\n"
        "az cosmosdb sql role assignment create \\\n"
        "  --account-name \"$COSMOS_ACCT\" --resource-group \"$RG\" \\\n"
        "  --role-definition-id \"00000000-0000-0000-0000-000000000002\" \\\n"
        "  --principal-id \"$MY_OID\" --scope \"/\""
    ))

    doc.add_heading("7.2 Add your public IP to the firewall", level=2)
    add_code_block(doc, (
        "MY_IP=$(curl -s https://api.ipify.org)\n"
        "az cosmosdb update --name \"$COSMOS_ACCT\" --resource-group \"$RG\" \\\n"
        "  --ip-range-filter \"...existing IPs...,$MY_IP\""
    ))

    doc.add_heading("7.3 Seed the four containers", level=2)
    doc.add_paragraph(
        "Adapt scripts/seed-dev-data.ts so it points at the live endpoint and uses "
        "DefaultAzureCredential instead of the dev key. Run it from the project root:"
    )
    add_code_block(doc, (
        "COSMOS_ENDPOINT=$(terraform output -raw cosmos_endpoint) \\\n"
        "COSMOS_DATABASE=dlp-database \\\n"
        "npx ts-node scripts/seed-prod-data.ts"
    ))

    add_callout(
        doc,
        "When done:",
        "Remove the temporary Cosmos role assignment and remove your IP from the "
        "firewall. Future Cosmos edits should go through the SWA admin path, not "
        "by repeatedly opening the firewall.",
    )

    add_check_items(doc, [
        "dlp-customers has at least one row.",
        "dlp-advisors, dlp-exemptions, dlp-encryption-exclusions populated to match your business rules.",
        "Your temporary Cosmos role + firewall opening have been revoked.",
    ])

    # ---------------- Phase 6: First deploy ----------------
    doc.add_heading("8. Phase 6 — First deploy through the pipeline", level=1)
    add_bullets(doc, [
        "Commit any local changes and push to main. The Azure DevOps pipeline starts automatically.",
        "BuildAndTest stage must pass with no lint or test failures (no continueOnError on lint).",
        "Infra stage runs terraform plan for visibility. Skim it.",
        "Deploy stage queues. Approve each environment (production-addin and production-api) when prompted.",
    ])

    doc.add_paragraph("After DeployApi finishes, smoke-test the API:")
    add_code_block(doc, (
        "curl -i \"https://$(terraform output -raw function_app_hostname)/api/config\"\n"
        "# Expect: HTTP/1.1 401   Missing Bearer token"
    ))

    add_callout(
        doc,
        "Security gate:",
        "A 401 with that exact body is the correct response — JWT validation is "
        "rejecting unauthenticated callers. A 200 here is a security bug; stop the "
        "rollout and investigate.",
    )

    doc.add_paragraph("Sanity-check the deployed manifest has no leftover placeholders:")
    add_code_block(doc, (
        "curl -s \"https://$(terraform output -raw static_web_app_hostname)/manifest.json\" \\\n"
        "  | grep -E \"REPLACE-WITH-|\\\\$\\\\{\" && echo \"❌ bad\" || echo \"✓ clean\""
    ))

    add_check_items(doc, [
        "/api/config returns 401 (not 200, not 500) without auth.",
        "SWA root URL returns 200 in a browser.",
        "Neither manifest contains REPLACE-WITH- or ${...} after deploy.",
    ])

    # ---------------- Phase 7: Pilot sideload ----------------
    doc.add_heading("9. Phase 7 — Pilot sideload (one test user)", level=1)
    doc.add_paragraph("Validate end-to-end on yourself before exposing the org.")
    add_bullets(doc, [
        "Download the XML manifest: curl -O https://<swa-hostname>/manifest-legacy.xml",
        "Open Outlook on the web  →  Settings (gear)  →  View all Outlook settings  →  General  →  Manage add-ins  →  My add-ins  →  Custom add-ins  →  Add a custom add-in  →  Add from file. Pick the downloaded XML.",
        "Open a new mail compose window — the DLP Guard ribbon button should appear.",
        "Click it. The taskpane should load and run the three checks against the current compose state.",
        "Open the browser DevTools (F12). Confirm no 401/403 errors on /api/config or /api/audit.",
        "Send a deliberately-blocking email (e.g., unencrypted .xlsx to an external recipient). Outlook should refuse with the DLP message.",
    ])

    add_check_items(doc, [
        "Ribbon button visible in compose.",
        "Taskpane renders without console errors.",
        "A known-bad email is blocked at Send time.",
        "An audit entry appears in Cosmos dlp-audit-log for the blocked send.",
    ])

    # ---------------- Phase 8: M365 rollout ----------------
    doc.add_heading("10. Phase 8 — Deploy to end users via Microsoft 365 Admin Center", level=1)
    add_bullets(doc, [
        "Open https://admin.microsoft.com/Adminportal/Home#/Settings/IntegratedApps.",
        "Upload custom apps  →  Office Add-in  →  Upload manifest file (.xml)  →  pick manifest-legacy.xml.",
        "Choose users: start with a pilot group of 5–10 users. After a week, expand.",
        "Microsoft 365 will show the permissions the add-in is requesting. Approve.",
        "Click Deploy. Distribution can take up to 12 hours on Outlook Desktop (minutes on Outlook Web).",
    ])

    add_callout(
        doc,
        "End-user impact:",
        "If you completed admin consent in Phase 3, users see no prompts — SSO is "
        "silent. A 'DLP Guard' button appears on the compose ribbon. When they click "
        "Send, the checks run invisibly; on a block, Outlook shows a dialog with the "
        "DLP message and a 'fix the issues' button.",
    )

    add_check_items(doc, [
        "Pilot users see the ribbon button.",
        "At least one pilot user has successfully sent a clean email.",
        "At least one pilot user has been blocked on a known-bad email.",
        "dlp-audit-log contains entries from pilot users.",
    ])

    # ---------------- End user experience ----------------
    doc.add_heading("11. What end users see and do", level=1)
    doc.add_paragraph(
        "Once the add-in is deployed via Integrated apps, end users do nothing — the "
        "add-in installs automatically. The next time they open Outlook (Web takes "
        "minutes; Desktop on Windows may take a few hours; Outlook for Mac requires a "
        "restart), they will notice:"
    )
    add_bullets(doc, [
        "A new 'DLP Guard' button on the Home ribbon of any New Email compose window.",
        "Clicking the button opens a sidebar pane that runs the three DLP checks on the current draft.",
        "If they click Send on a draft that fails the checks, Outlook stops the send and shows a dialog explaining what to fix (in Hebrew, e.g. 'קבצים לא מוצפנים: payroll.xlsx').",
        "If the draft is clean, Send proceeds normally — there is no extra delay or extra dialog.",
    ])

    doc.add_paragraph(
        "If a user tries to send while the DLP service is unreachable (network blip, "
        "Azure outage), the add-in fails open — the email is sent — but emits an "
        "audit event with action = DLP_UNAVAILABLE so operators can see coverage gaps."
    )

    doc.add_heading("Troubleshooting (end-user-facing issues)", level=2)
    add_table(
        doc,
        ["Symptom", "Probable cause", "Fix"],
        [
            ["Ribbon button missing", "Manifest hasn't propagated yet", "Wait up to 12 hours; restart Outlook; ask the user to try Outlook on the web first"],
            ["'SSO token failed code=13007'", "Office host doesn't support SSO (e.g., Outlook mobile)", "Expected — DLP fails open on these hosts; check Microsoft's Office support matrix"],
            ["DLP blocks every email", "Cosmos dlp-customers / dlp-encryption-exclusions empty or misconfigured", "Re-seed via Phase 5 procedure"],
            ["DLP blocks emails it shouldn't", "Magic-byte detection cannot read the attachment (mobile, very large file)", "The block message will say 'cannot verify' — user should try from a desktop client"],
            ["Send works but no audit entries appear", "/api/audit returning 401", "Check the Functions ALLOWED_AUDIENCE matches the Entra app identifier URI exactly"],
        ],
    )

    # ---------------- Operations ----------------
    doc.add_heading("12. Day-2 operations", level=1)

    doc.add_heading("Updates", level=2)
    add_bullets(doc, [
        "Code changes  →  merge to main  →  pipeline auto-builds  →  approve the deploy.",
        "Manifest changes (rare)  →  bump the version field in manifest-legacy.xml first, then redistribute via Integrated apps.",
        "Cosmos data (customers, exemptions, etc.)  →  use the temporary-role pattern from Phase 5, or build a small admin SPA.",
    ])

    doc.add_heading("Monitoring", level=2)
    add_bullets(doc, [
        "Application Insights resource 'nextage-dlp-ai' — failed requests, latency.",
        "Cosmos dlp-audit-log — search for action == 'DLP_UNAVAILABLE' to see fail-open events.",
        "Static Web Apps logs — diagnose taskpane HTML / JS issues.",
    ])

    doc.add_heading("Rotating the HMAC key", level=2)
    doc.add_paragraph(
        "Rotation is destructive for historical hash comparability. Only rotate "
        "during an incident response:"
    )
    add_bullets(doc, [
        "openssl rand -hex 32  →  new key.",
        "Update audit_hmac_key in terraform.tfvars and apply.",
        "Document the rotation date in your incident log; new and old hashes cannot be correlated.",
    ])

    doc.add_heading("Decommissioning", level=2)
    add_code_block(doc, "cd terraform && terraform destroy")
    doc.add_paragraph(
        "Then remove the integrated app from the Microsoft 365 Admin Center, and "
        "double-check the Entra app registration is gone in Entra ID."
    )

    # ---------------- Reference ----------------
    doc.add_heading("13. Variable source-of-truth reference", level=1)
    add_table(
        doc,
        ["Where it lives", "What it is", "Source"],
        [
            ["MANIFEST_GUID  →  manifest <Id>", "Stable add-in identifier", "terraform.tfvars (set once)"],
            ["ENTRA_APP_CLIENT_ID  →  manifest webApplicationInfo.id", "Entra app's client ID", "terraform output entra_app_client_id"],
            ["SWA_HOSTNAME  →  manifest URLs and validDomains", "Static Web App default hostname", "terraform output manifest_tokens.SWA_HOSTNAME"],
            ["FUNCTIONS_HOSTNAME  →  manifest validDomains / AppDomains", "Function App default hostname", "terraform output manifest_tokens.FUNCTIONS_HOSTNAME"],
            ["ALLOWED_AUDIENCE  (Function app setting)", "JWT 'aud' accepted by API", "terraform output entra_app_identifier_uri"],
            ["AZURE_TENANT_ID  (Function app setting)", "JWT 'tid' accepted by API", "terraform.tfvars"],
            ["AUDIT_HMAC_KEY  (Function app setting)", "PII redaction key", "terraform.tfvars"],
        ],
    )

    # Done
    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    build()
